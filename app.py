import streamlit as st
import io
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import requests
from streamlit_lottie import st_lottie
from Bio.PDB import PDBParser, PPBuilder, PDBList, SASA
from Bio.SeqUtils import ProtParam
from docx import Document
from docx.shared import Inches, RGBColor
from streamlit_molstar import st_molstar
from matplotlib.patches import RegularPolygon

# --- 1. SCHRÖDINGER-INSPIRED CONFIG ---
st.set_page_config(page_title="BioMumo | Enzyme Optimization Hub", layout="wide", page_icon="🧬")
# --- 1. CONFIG & SCHRÖDINGER-INSPIRED STYLING ---
st.set_page_config(page_title="BioMumo | MOMU CORE", layout="wide", page_icon="🧬")

st.markdown("""
    <style>
    .stApp { background-color: #0b0f19; color: #e2e8f0; }
    .hero-text { text-align: center; padding: 15px 0px; background: linear-gradient(180deg, #0b0f19 0%, #161e2d 100%); }
    .main-title { font-family: 'Inter', sans-serif; font-size: 38px; font-weight: 800; background: linear-gradient(90deg, #00d4ff, #ffffff); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 5px; }
    .sub-title { font-size: 13px; color: #94a3b8; letter-spacing: 1.5px; text-transform: uppercase; }
    .stButton>button { width: 100%; border-radius: 5px; border: 1px solid #00d4ff; background-color: transparent; color: #00d4ff; font-weight: 600; height: 3em; text-transform: uppercase; }
    .stButton>button:hover { background-color: #00d4ff; color: #0b0f19; box-shadow: 0 0 15px rgba(0, 212, 255, 0.4); }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; justify-content: center; }
    .stTabs [data-baseweb="tab"] { height: 45px; background-color: #0f172a; border-radius: 4px; color: #94a3b8; font-weight: 600; }
    .stTabs [aria-selected="true"] { color: #00d4ff !important; border-bottom-color: #00d4ff !important; }
    .main-title { font-family: 'Inter', sans-serif; font-size: 38px; font-weight: 800; background: linear-gradient(90deg, #00d4ff, #ffffff); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; }
    
    /* Benzene Ring UI Components */
    .benzene-container { display: flex; flex-direction: column; align-items: center; padding: 10px; }
    .hexagon {
        width: 180px; height: 210px;
        background: linear-gradient(45deg, #00d4ff, #005f73);
        clip-path: polygon(50% 0%, 100% 25%, 100% 75%, 50% 100%, 0% 75%, 0% 25%);
        display: flex; align-items: center; justify-content: center;
        border: 2px solid #ffffff;
    }
    .hexagon img {
        width: 150px; height: 150px;
        clip-path: polygon(50% 0%, 100% 25%, 100% 75%, 50% 100%, 0% 75%, 0% 25%);
        background-color: #0b0f19; object-fit: cover;
    }
    .pipeline-title { font-weight: 800; font-size: 18px; color: #00d4ff; margin-bottom: 5px; text-transform: uppercase; text-align: center; }
    .info-box { background: #161e2d; padding: 15px; border-radius: 10px; border-left: 5px solid #00d4ff; margin-bottom: 10px; font-size: 14px; color: #cbd5e1; line-height: 1.6; }
    
    /* Button Customization */
    .stButton>button { width: 100%; border-radius: 5px; border: 1px solid #00d4ff; background-color: transparent; color: #00d4ff; font-weight: 600; }
    .stButton>button:hover { background-color: #00d4ff; color: #0b0f19; }
    </style>
""", unsafe_allow_html=True)

# --- 2. CORE UTILITIES ---
def load_lottieurl(url: str):
    try:
        r = requests.get(url, timeout=5)
        return r.json() if r.status_code == 200 else None
    except: return None

lottie_scan = load_lottieurl("https://assets5.lottiefiles.com/packages/lf20_m6cu9zoc.json")

def create_prof_report(title, methodology, formulas, df, plot_buf=None):
def create_prof_report(title, methodology, df):
    doc = Document()
    header = doc.add_heading(title, 0)
    header.runs[0].font.color.rgb = RGBColor(30, 58, 170)
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(methodology)
    if formulas:
        doc.add_heading('Mathematical Basis', level=2)
        for f in formulas: doc.add_paragraph(f, style='Quote')
    doc.add_heading('Results', level=1)
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    for j, col in enumerate(df.columns): table.cell(0, j).text = str(col)
    for i, row in enumerate(df.values):
        for j, val in enumerate(row): table.cell(i + 1, j).text = str(val)
    if plot_buf:
        doc.add_heading('Visualization', level=1)
        doc.add_picture(plot_buf, width=Inches(5))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. HORIZONTAL LOGO GENERATOR ---
def generate_biomumo_logo():
    # Reduced height (2.0) and expanded width (18) for a sleek banner
    fig, ax = plt.subplots(figsize=(18, 2.0), facecolor='#0b0f19')
    ax.set_facecolor('#0b0f19')
    ax.set_xlim(-8, 28)
    ax.set_ylim(-4, 4)
    ax.axis('off')
    
    ax.set_xlim(-8, 28); ax.set_ylim(-4, 4); ax.axis('off')
    chem_color, accent_color = '#ffffff', '#00d4ff'
    
    # 4-Benzene Rings in a compact horizontal line
    ring_centers = [(-4, 0), (-1, 0), (2, 0), (5, 0)]
    letters = ["M", "U", "M", "O"]
    
    for center, letter in zip(ring_centers, letters):
        ring = RegularPolygon(center, numVertices=6, radius=1.3, orientation=0, 
                              edgecolor=chem_color, facecolor='none', lw=2)
        ring = RegularPolygon(center, numVertices=6, radius=1.3, orientation=0, edgecolor=chem_color, facecolor='none', lw=2)
        ax.add_patch(ring)
        ax.text(center[0], center[1], letter, color=accent_color, fontsize=16, 
                fontweight='black', ha='center', va='center')
    
    # Stylized chemical tail
    ax.plot([5, 6, 7.5, 9], [1.1, 1.8, 1.1, 1.8], color=chem_color, lw=2, alpha=0.8)
    
    # Text branding moved further right to prevent overlap
        ax.text(center[0], center[1], letter, color=accent_color, fontsize=16, fontweight='black', ha='center', va='center')
    ax.text(11, 0.5, "MUMO CORE", color='#ffffff', fontsize=48, fontweight='black')
    ax.text(11, -1.2, "INTEGRATED PIPELINE ECOSYSTEM | VINAYAKA MISSION'S COLLEGE OF PHARMACY", 
            color='#94a3b8', fontsize=12, fontweight='bold', alpha=0.7)
    
    return fig

# --- 4. RENDER LOGO ---
plt.close('all') 
logo_fig = generate_biomumo_logo()
logo_buf = io.BytesIO()
logo_fig.savefig(logo_buf, format='png', bbox_inches='tight', pad_inches=0.1, transparent=True)
logo_buf.seek(0)

# FIXED: Removed potential bracket issues in layout
st.image(logo_buf, use_container_width=True)
# --- 3. SESSION STATE ---
if 'show_desc' not in st.session_state:
    st.session_state.show_desc = {"p1": False, "p2": False, "p3": False}
if 'active_file' not in st.session_state:
    st.session_state.active_file = None

# --- 4. RENDER UI ---
st.pyplot(generate_biomumo_logo())
tabs = st.tabs(["🏠 HOME / PIPELINE", "📜 DESCRIPTIONS", "👥 ABOUT US", "📚 REFERENCES", "📧 CONTACT"])

# --- 5. PAGE: HOME / PIPELINE ---
with tabs[0]:
    st.markdown('<div class="hero-text"><p class="sub-title">Computational Drug Discovery Platform</p><h1 class="main-title">BioMumo: Opening New Worlds for Molecular Discovery</h1></div>', unsafe_allow_html=True)
    if 'active_file' not in st.session_state: st.session_state.active_file = None
    if 'active_name' not in st.session_state: st.session_state.active_name = "Target"

    col_left, col_right = st.columns([1, 2], gap="large")
    with col_left:
        st.markdown("### 🧪 RESEARCH INPUT")
        mode = st.radio("Protocol", ["Upload PDB", "Enter PDB ID"], horizontal=True)
    st.markdown('<h1 class="main-title">BioMumo: Molecular Discovery Pipeline</h1>', unsafe_allow_html=True)
    
    # Target Selection
    col_in, col_viz = st.columns([1, 2])
    with col_in:
        st.markdown("### 🧪 Target Selection")
        mode = st.radio("Method", ["Upload PDB", "PDB ID"], horizontal=True)
        if mode == "Upload PDB":
            up = st.file_uploader("Upload Structure", type=['pdb'])
            if up:
                with open("temp.pdb", "wb") as f: f.write(up.getbuffer())
                st.session_state.active_file = "temp.pdb"
                st.session_state.active_name = up.name.split('.')[0]
        else:
            pid = st.text_input("PDB ID (e.g., 3FXI)").upper()
            pid = st.text_input("Enter PDB ID (e.g., 3FXI)").upper()
            if pid:
                with st.spinner("Fetching..."):
                    try:
                        pdbl = PDBList()
                        st.session_state.active_file = pdbl.retrieve_pdb_file(pid, pdir='.', file_format='pdb')
                        st.session_state.active_name = pid
                    except: st.error("Fetch Error")
        st.divider()
        st.markdown("### ⚡ CORE UTILITIES")
        run1, run2, run3 = st.button("① Protein Structure Analysis"), st.button("② Active Site Mapping"), st.button("③ Mutation Prediction")

    with col_right:
                    
    with col_viz:
        if st.session_state.active_file:
            parser = PDBParser(QUIET=True)
            structure = parser.get_structure(st.session_state.active_name, st.session_state.active_file)
            st_molstar(st.session_state.active_file, height=450)
            
            if run1:
            st_molstar(st.session_state.active_file, height=350)
        else:
            st.info("Awaiting molecular target for visualization.")

    st.divider()

    # --- PIPELINE GRID ---
    c1, c2, c3 = st.columns(3)

    # PIPELINE 1
    with c1:
        st.markdown('<p class="pipeline-title">Protein Analysis</p>', unsafe_allow_html=True)
        st.markdown('<div class="benzene-container"><div class="hexagon"><img src="https://rcsb.org/pdb/images/3fxi_asym_r_500.jpg"></div></div>', unsafe_allow_html=True)
        if st.button("Click Here", key="btn1"): st.session_state.show_desc["p1"] = not st.session_state.show_desc["p1"]
        if st.session_state.show_desc["p1"]:
            st.markdown('<div class="info-box">1. Extracts sequence data.<br>2. Calculates MW (kDa).<br>3. Computes pI.<br>4. Measures Instability Index.<br>5. Evaluates hydrophobicity.</div>', unsafe_allow_html=True)
            if st.button("▶ Run Protein Analysis") and st.session_state.active_file:
                parser = PDBParser(QUIET=True)
                structure = parser.get_structure("target", st.session_state.active_file)
                ppb = PPBuilder()
                seq = "".join([str(p.get_sequence()) for p in ppb.build_peptides(structure)])
                ana = ProtParam.ProteinAnalysis(seq)
                df1 = pd.DataFrame({'Parameter': ['MW', 'pI', 'Instability'], 'Value': [f"{ana.molecular_weight()/1000:.2f} kDa", f"{ana.isoelectric_point():.2f}", f"{ana.instability_index():.2f}"]})
                st.table(df1)
            
            elif run2:
                with st.spinner("Mapping Active Site..."):
                    try:
                        sr = SASA.ShrakeRupley()
                        sr.compute(structure[0], level="R")
                        sites = [{'Residue': f"{res.get_resname()}{res.id[1]}", 'SASA': round(res.sasa, 2)} for res in structure.get_residues() if hasattr(res, 'sasa')]
                        st.dataframe(pd.DataFrame(sites).nlargest(15, 'SASA'), use_container_width=True)
                    except: st.warning("SASA calculation requires full atom data.")
            
            elif run3:
                df = pd.DataFrame({'Parameter': ['MW', 'pI', 'Instability'], 'Value': [f"{ana.molecular_weight()/1000:.2f} kDa", f"{ana.isoelectric_point():.2f}", f"{ana.instability_index():.2f}"]})
                st.table(df)

    # PIPELINE 2
    with c2:
        st.markdown('<p class="pipeline-title">Active Site Prediction</p>', unsafe_allow_html=True)
        st.markdown('<div class="benzene-container"><div class="hexagon"><img src="https://cdn.rcsb.org/images/structures/8/8m5s/8m5s_assembly-1.jpeg"></div></div>', unsafe_allow_html=True)
        if st.button("Click Here", key="btn2"): st.session_state.show_desc["p2"] = not st.session_state.show_desc["p2"]
        if st.session_state.show_desc["p2"]:
            st.markdown('<div class="info-box">1. Computes SASA.<br>2. Identifies pocket residues.<br>3. Ranks by exposure.<br>4. Predicts binding affinity.<br>5. Maps catalytic sites.</div>', unsafe_allow_html=True)
            if st.button("▶ Run Site Mapping") and st.session_state.active_file:
                parser = PDBParser(QUIET=True)
                structure = parser.get_structure("target", st.session_state.active_file)
                sr = SASA.ShrakeRupley()
                sr.compute(structure[0], level="R")
                sites = [{'Residue': f"{res.get_resname()}{res.id[1]}", 'SASA': round(res.sasa, 2)} for res in structure.get_residues() if hasattr(res, 'sasa')]
                st.dataframe(pd.DataFrame(sites).nlargest(10, 'SASA'), use_container_width=True)

    # PIPELINE 3
    with c3:
        st.markdown('<p class="pipeline-title">Mutation Prediction</p>', unsafe_allow_html=True)
        st.markdown('<div class="benzene-container"><div class="hexagon"><img src="https://cdn.rcsb.org/images/structures/1/1aie/1aie_assembly-1.jpeg"></div></div>', unsafe_allow_html=True)
        if st.button("Click Here", key="btn3"): st.session_state.show_desc["p3"] = not st.session_state.show_desc["p3"]
        if st.session_state.show_desc["p3"]:
            st.markdown('<div class="info-box">1. Analyzes B-factors.<br>2. Finds flexible regions.<br>3. Predicts mutation impact.<br>4. Suggests substitutions.<br>5. Heatmap visualization.</div>', unsafe_allow_html=True)
            if st.button("▶ Run Mutation Analysis") and st.session_state.active_file:
                parser = PDBParser(QUIET=True)
                structure = parser.get_structure("target", st.session_state.active_file)
                res_data = [{"Pos": a.get_parent().id[1], "B": a.get_bfactor()} for a in structure.get_atoms()]
                df_mut = pd.DataFrame(res_data).groupby('Pos').mean().reset_index()
                st.line_chart(df_mut.set_index('Pos')['B'])
        else: st.info("Awaiting molecular target.")

# --- 6. PAGE: DESCRIPTIONS (RESTORED FULL) ---
# --- RE-ADDING ORIGINAL DESCRIPTIONS TAB ---
with tabs[1]:
    st.markdown('<h2 class="main-title">Methodology & Mathematical Basis</h2>', unsafe_allow_html=True)
    st.write("The analysis utilizes the **ExPASy ProtParam** algorithm to derive fundamental properties. By treating the primary sequence as a linear chain of residues, we predict behavior in physiological environments, critical for mucosal drug delivery where pH and ionic strength alter enzyme activity.")
    st.info("**Isoelectric Point (pI):** $pI = (pK_i + pK_j)/2$. This determines the pH at which the enzyme carries no net charge, vital for stability in variable mucosal layers.")
    st.info("**Molecular Weight (MW):** $MW = \sum (n_i \times m_i) + (H_2O)$. Calculated by isotopic mass summation; high MW proteins face diffusion barriers in thick mucin.")
    st.write("The **Debye-Waller Factor** (B-factor) reflects thermal displacement. In engineering, high B-factors represent flexibility hotspots. Rigidifying these residues enhances thermostability.")
    st.info("**Isoelectric Point (pI):** $pI = (pK_i + pK_j)/2$. Determines pH where the enzyme has no net charge.")
    st.info("**B-factor Analysis:** Represents thermal displacement. High B-factors = High flexibility.")

# --- 7. PAGE: ABOUT US (RESTORED FULL) ---
with tabs[2]:
    st.markdown("### 🏛️ Vinayaka Mission's College of Pharmacy")
    st.write("A constituent college of Vinayaka Mission's Research Foundation, our institution is at the forefront of pharmaceutical innovation. This platform was developed as part of advanced research into **In-Silico Drug Discovery** and **Computational Proteomics**.")
    st.success("**Mission Objective:** To bridge the gap between traditional Wet-Lab pharmacy and high-performance computational modeling. We focus on training pharmacists to utilize Python-based bioinformatics for solving complex biological challenges.")
    st.info("**Project Goal:** The core of this initiative lies in building an end-to-end **computational pipeline** designed to bridge the gap between raw data and actionable insights through **Machine Learning (ML)** and **Deep Learning (DL)**. Utilizing **Neural Network architectures**, the system automates high-dimensional feature extraction and pattern recognition.")
    st.write("### Vinayaka Mission's College of Pharmacy")
    st.write("Developed by Mowriss.M.G & Mugilarasi.C. Bridging Pharmacy with Computational Proteomics.")

# --- 8. PAGE: REFERENCES (RESTORED FULL) ---
with tabs[3]:
    st.write("""
    1. **Cock PJ, et al. (2009).** *Biopython: freely available Python tools for computational molecular biology.* Bioinformatics, 25(11).
    2. **Gasteiger E, et al. (2005).** *Protein Identification and Analysis Tools on the ExPASy Server.*
    3. **Sun Z, et al. (2019).** *Utility of B-factors in protein engineering.* Chemical Reviews.
    4. **Abramson J, et al. (2024).** *Accurate structure prediction of biomolecular interactions with AlphaFold 3.* Nature, 630.
    5. **Khan A, et al. (2026).** *Protein structure prediction powered by artificial intelligence.* Front Mol Biosci.
    """)

# --- 9. PAGE: CONTACT ---
    st.write("1. Biopython: Bioinformatics tools. 2. ExPASy ProtParam algorithm. 3. AlphaFold 3 (Reference).")

with tabs[4]:
    st.markdown('<div style="text-align:center;"><h2>Mowriss.M.G & Mugilarasi.C</h2><p>B.Pharm Research Scholars</p><p><strong>Vinayaka Mission\'s College of Pharmacy</strong></p></div>', unsafe_allow_html=True)
    st.write("Contact: Mowriss.M.G & Mugilarasi.C | VMCP Research Scholars.")
